from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_COLOR_INDEX
import os
import re
from typing import Dict, List, Any

class DocumentParser:
    def __init__(self):
        self.document_types = {
            'articles of association': 'Articles of Association',
            'memorandum of association': 'Memorandum of Association',
            'board resolution': 'Board Resolution',
            'shareholder resolution': 'Shareholder Resolution',
            'incorporation application': 'Incorporation Application Form',
            'ubo declaration': 'UBO Declaration Form',
            'register of members': 'Register of Members and Directors',
            'employment contract': 'Employment Contract',
            'license application': 'License Application'
        }
    
    def parse_document(self, file_path: str) -> Dict[str, Any]:
        """Parse Word document and extract structured information"""
        try:
            doc = Document(file_path)
            
            # Extract basic document info
            doc_data = {
                'file_path': file_path,
                'document_type': self._identify_document_type(doc),
                'content': self._extract_text_content(doc),
                'paragraphs': self._extract_paragraphs(doc),
                'tables': self._extract_tables(doc),
                'word_count': len(self._extract_text_content(doc).split()),
                'structure': self._analyze_document_structure(doc),
                'metadata': self._extract_metadata(doc)
            }
            
            return doc_data
            
        except Exception as e:
            raise Exception(f"Error parsing document: {str(e)}")
    
    def _identify_document_type(self, doc: Document) -> str:
        """Identify document type based on content"""
        content = self._extract_text_content(doc).lower()
        
        # Check for specific keywords and patterns
        for keywords, doc_type in self.document_types.items():
            if keywords in content:
                return doc_type
        
        # Additional pattern matching
        if 'articles' in content and 'association' in content:
            return 'Articles of Association'
        elif 'memorandum' in content and 'association' in content:
            return 'Memorandum of Association'
        elif 'resolution' in content and 'board' in content:
            return 'Board Resolution'
        elif 'shareholders' in content and 'resolution' in content:
            return 'Shareholder Resolution'
        elif 'ultimate beneficial owner' in content or 'ubo' in content:
            return 'UBO Declaration Form'
        elif 'register' in content and ('members' in content or 'directors' in content):
            return 'Register of Members and Directors'
        elif 'employment' in content and 'contract' in content:
            return 'Employment Contract'
        elif 'license' in content and 'application' in content:
            return 'License Application'
        
        return 'Unknown Document Type'
    
    def _extract_text_content(self, doc: Document) -> str:
        """Extract all text content from document"""
        content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text.strip())
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        content.append(cell.text.strip())
        
        return '\n'.join(content)
    
    def _extract_paragraphs(self, doc: Document) -> List[Dict[str, Any]]:
        """Extract paragraphs with formatting information"""
        paragraphs = []
        
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                para_data = {
                    'index': i,
                    'text': paragraph.text.strip(),
                    'style': paragraph.style.name if paragraph.style else 'Normal',
                    'is_heading': self._is_heading(paragraph),
                    'has_numbering': self._has_numbering(paragraph),
                    'formatting': self._get_paragraph_formatting(paragraph)
                }
                paragraphs.append(para_data)
        
        return paragraphs
    
    def _extract_tables(self, doc: Document) -> List[Dict[str, Any]]:
        """Extract tables with content"""
        tables = []
        
        for i, table in enumerate(doc.tables):
            table_data = {
                'index': i,
                'rows': len(table.rows),
                'columns': len(table.columns),
                'content': []
            }
            
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data['content'].append(row_data)
            
            tables.append(table_data)
        
        return tables
    
    def _analyze_document_structure(self, doc: Document) -> Dict[str, Any]:
        """Analyze document structure and organization"""
        structure = {
            'total_paragraphs': len(doc.paragraphs),
            'total_tables': len(doc.tables),
            'headings': [],
            'sections': [],
            'has_toc': False,
            'has_signature_section': False,
            'has_date_section': False
        }
        
        content = self._extract_text_content(doc).lower()
        
        # Check for signature section
        signature_keywords = ['signature', 'signed', 'witness', 'executed']
        structure['has_signature_section'] = any(keyword in content for keyword in signature_keywords)
        
        # Check for date section
        date_patterns = [r'\d{1,2}[/-]\d{1,2}[/-]\d{4}', r'\d{1,2}\s+\w+\s+\d{4}']
        structure['has_date_section'] = any(re.search(pattern, content) for pattern in date_patterns)
        
        # Extract headings
        for i, paragraph in enumerate(doc.paragraphs):
            if self._is_heading(paragraph):
                structure['headings'].append({
                    'index': i,
                    'text': paragraph.text.strip(),
                    'level': self._get_heading_level(paragraph)
                })
        
        return structure
    
    def _extract_metadata(self, doc: Document) -> Dict[str, Any]:
        """Extract document metadata"""
        core_props = doc.core_properties
        
        metadata = {
            'title': core_props.title or '',
            'author': core_props.author or '',
            'subject': core_props.subject or '',
            'created': core_props.created.isoformat() if core_props.created else None,
            'modified': core_props.modified.isoformat() if core_props.modified else None,
            'last_modified_by': core_props.last_modified_by or '',
            'revision': core_props.revision or 0,
            'version': core_props.version or ''
        }
        
        return metadata
    
    def _is_heading(self, paragraph) -> bool:
        """Check if paragraph is a heading"""
        if paragraph.style:
            style_name = paragraph.style.name.lower()
            return 'heading' in style_name or 'title' in style_name
        return False
    
    def _get_heading_level(self, paragraph) -> int:
        """Get heading level"""
        if paragraph.style:
            style_name = paragraph.style.name.lower()
            if 'heading' in style_name:
                # Extract number from "Heading 1", "Heading 2", etc.
                level_match = re.search(r'heading\s+(\d+)', style_name)
                if level_match:
                    return int(level_match.group(1))
        return 1
    
    def _has_numbering(self, paragraph) -> bool:
        """Check if paragraph has numbering"""
        text = paragraph.text.strip()
        # Check for common numbering patterns
        patterns = [
            r'^\d+\.',  # 1., 2., 3.
            r'^\d+\)',  # 1), 2), 3)
            r'^\(\d+\)',  # (1), (2), (3)
            r'^[a-z]\.',  # a., b., c.
            r'^[A-Z]\.',  # A., B., C.
            r'^\([a-z]\)',  # (a), (b), (c)
            r'^\([A-Z]\)'   # (A), (B), (C)
        ]
        
        return any(re.match(pattern, text) for pattern in patterns)
    
    def _get_paragraph_formatting(self, paragraph) -> Dict[str, Any]:
        """Get paragraph formatting information"""
        formatting = {
            'bold': False,
            'italic': False,
            'underline': False,
            'font_size': None,
            'font_name': None,
            'alignment': None
        }
        
        if paragraph.runs:
            # Check first run for formatting
            first_run = paragraph.runs[0]
            formatting['bold'] = first_run.bold or False
            formatting['italic'] = first_run.italic or False
            formatting['underline'] = first_run.underline or False
            
            if first_run.font.size:
                formatting['font_size'] = first_run.font.size.pt
            if first_run.font.name:
                formatting['font_name'] = first_run.font.name
        
        if paragraph.paragraph_format.alignment:
            formatting['alignment'] = str(paragraph.paragraph_format.alignment)
        
        return formatting
