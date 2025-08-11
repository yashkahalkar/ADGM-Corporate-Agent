import streamlit as st
import os
import time
from pathlib import Path
from pinecone import Pinecone, ServerlessSpec
from sentence_transformers import SentenceTransformer  
from models.gemini_client import GeminiClient
import numpy as np
from typing import List, Dict, Any
import json
import requests
from bs4 import BeautifulSoup
from docx import Document
import urllib.request


class RAGEngine:
    def __init__(self):
        print("Initializing Production ADGM RAG Engine...")
        
        # Streamlit secrets integration for Pinecone
        try:
            api_key = st.secrets['PINECONE_API_KEY']
        except (KeyError, FileNotFoundError, AttributeError):
            api_key = os.getenv('PINECONE_API_KEY')
            if not api_key:
                raise ValueError("PINECONE_API_KEY not found in secrets or environment")
        
        self.pc = Pinecone(api_key=api_key)
        
        try:
            self.index_name = st.secrets.get('PINECONE_INDEX_NAME', 'adgm-legal-docs')
        except (KeyError, FileNotFoundError, AttributeError):
            self.index_name = os.getenv('PINECONE_INDEX_NAME', 'adgm-legal-docs')

class RAGEngine:
    def __init__(self):
        print("Initializing Production ADGM RAG Engine...")
        
        # Initialize Pinecone
        try:
            api_key = st.secrets['PINECONE_API_KEY']
        except (KeyError, FileNotFoundError, AttributeError):
            api_key = os.getenv('PINECONE_API_KEY')
            if not api_key:
                raise ValueError("PINECONE_API_KEY not found in Streamlit secrets or environment variables")
        
        try:
            self.index_name = st.secrets.get('PINECONE_INDEX_NAME', 'adgm-legal-docs')
        except (KeyError, FileNotFoundError, AttributeError):
            self.index_name = os.getenv('PINECONE_INDEX_NAME', 'adgm-legal-docs')
        
        # Create documents directory
        self.docs_dir = Path("official_adgm_documents")
        self.docs_dir.mkdir(exist_ok=True)
        
        print("Loading embedding model...")
        self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
        self.gemini_client = GeminiClient()
        
        # Setup index
        self._setup_index()
        time.sleep(5)
        
        # Connect to index
        self.index = self.pc.Index(self.index_name)
        
        # Download official documents and populate knowledge base
        if self._is_index_empty():
            print("Downloading official ADGM documents and populating knowledge base...")
            self._download_official_documents()
            self._populate_from_official_documents()
        
        print("Production ADGM RAG Engine initialized successfully!")
    
    def _setup_index(self):
        """Setup Pinecone index"""
        try:
            existing_indexes = [index.name for index in self.pc.list_indexes()]
            print(f"Existing indexes: {existing_indexes}")
            
            if self.index_name not in existing_indexes:
                print(f"Creating production index: {self.index_name}")
                self.pc.create_index(
                    name=self.index_name,
                    dimension=384,
                    metric='cosine',
                    spec=ServerlessSpec(cloud='aws', region='us-east-1')
                )
                print(f"Production index {self.index_name} created!")
            else:
                print(f"Using existing production index: {self.index_name}")
        except Exception as e:
            print(f"Error setting up index: {str(e)}")
            raise
    
    def _is_index_empty(self) -> bool:
        """Check if index is empty"""
        try:
            stats = self.index.describe_index_stats()
            return stats['total_vector_count'] == 0
        except Exception as e:
            print(f"Error checking index stats: {str(e)}")
            return True
    
    def _download_official_documents(self):
        """Download official ADGM documents"""
        print("Downloading official ADGM templates...")
        
        # Official ADGM document URLs
        official_documents = {
            'model_articles_shares.docx': 'https://assets.adgm.com/download/assets/adgm-ra-model-articles-private-company-limited-by-shares.docx/015402647f0111ef91cdea7ac70a8286',
            'model_articles_guarantee.docx': 'https://assets.adgm.com/download/assets/adgm-ra-model-articles-private-company-limited-by-guarantee.docx/e6d3adc05b1711ef9f15a617eb0b5f27',
            'board_resolution_business_activities.docx': 'https://assets.adgm.com/download/assets/Templates_BoardReso_BusinessActivitiesChange-v1-20220107.docx/c4866d8c5b0011efbb2c8ea8406205f9',
            'board_resolution_trade_name.docx': 'https://assets.adgm.com/download/assets/Templates_BoardReso_TradeName-v1-20220107.docx/90dd74085b0511ef98b1fe647dc54e16',
            'board_resolution_directors_resignation.docx': 'https://assets.adgm.com/download/assets/Templates_BoardReso_DirectorsResignation-v1-20220107.docx/86ea7d805b0311efabc1da675695d30e',
            'shareholders_resolution_general.docx': 'https://assets.adgm.com/download/assets/UNOFFICIAL---Template-Shareholders-Resolution.docx/44fe98a85d4611efac830241e0190a99',
            'incorporation_resolution_individual.docx': 'https://assets.adgm.com/download/assets/adgm-ra-resolution-single-individual-shareholder-LTD-incorporation-v2.docx/f160dbe06c3911efa22adaed20215b4a',
            'incorporation_resolution_corporate.docx': 'https://assets.adgm.com/download/assets/adgm-ra-resolution-incorporation-ltd-corporate-shareholder-v2.docx/983698446c3811efbf7cfe7434582912'
        }
        
        for filename, url in official_documents.items():
            file_path = self.docs_dir / filename
            
            # Only download if file doesn't exist
            if not file_path.exists():
                try:
                    print(f"Downloading {filename}...")
                    headers = {
                        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
                    }
                    
                    req = urllib.request.Request(url, headers=headers)
                    with urllib.request.urlopen(req) as response:
                        with open(file_path, 'wb') as f:
                            f.write(response.read())
                    
                    print(f"✅ Downloaded {filename}")
                except Exception as e:
                    print(f"❌ Failed to download {filename}: {str(e)}")
            else:
                print(f"✅ {filename} already exists")
    
    def _populate_from_official_documents(self):
        """Populate knowledge base from downloaded official DOCX files"""
        print("Reading official ADGM documents and creating embeddings...")
        
        # Document metadata mapping
        document_metadata = {
            'model_articles_shares.docx': {
                'content_type': 'official_template',
                'category': 'incorporation',
                'document_type': 'Model Articles - Private Company Limited by Shares',
                'source': 'ADGM Registration Authority'
            },
            'model_articles_guarantee.docx': {
                'content_type': 'official_template',
                'category': 'incorporation',
                'document_type': 'Model Articles - Private Company Limited by Guarantee',
                'source': 'ADGM Registration Authority'
            },
            'board_resolution_business_activities.docx': {
                'content_type': 'official_template',
                'category': 'corporate_governance',
                'document_type': 'Board Resolution - Business Activities Change',
                'source': 'ADGM Registration Authority'
            },
            'board_resolution_trade_name.docx': {
                'content_type': 'official_template',
                'category': 'corporate_governance',
                'document_type': 'Board Resolution - Trade Name',
                'source': 'ADGM Registration Authority'
            },
            'board_resolution_directors_resignation.docx': {
                'content_type': 'official_template',
                'category': 'corporate_governance',
                'document_type': 'Board Resolution - Directors Resignation',
                'source': 'ADGM Registration Authority'
            },
            'shareholders_resolution_general.docx': {
                'content_type': 'official_template',
                'category': 'corporate_governance',
                'document_type': 'Shareholders Resolution - General',
                'source': 'ADGM Registration Authority'
            },
            'incorporation_resolution_individual.docx': {
                'content_type': 'official_template',
                'category': 'incorporation',
                'document_type': 'Incorporation Resolution - Individual Shareholder',
                'source': 'ADGM Registration Authority'
            },
            'incorporation_resolution_corporate.docx': {
                'content_type': 'official_template',
                'category': 'incorporation',
                'document_type': 'Incorporation Resolution - Corporate Shareholder',
                'source': 'ADGM Registration Authority'
            }
        }
        
        vectors = []
        
        for filename, metadata in document_metadata.items():
            file_path = self.docs_dir / filename
            
            if file_path.exists():
                try:
                    print(f"Processing {filename}...")
                    
                    # Read DOCX file
                    content = self._read_docx_file(file_path)
                    
                    if content:
                        # Split into chunks
                        chunks = self._chunk_content(content)
                        print(f"Created {len(chunks)} chunks from {filename}")
                        
                        for i, chunk in enumerate(chunks):
                            # Create embedding
                            embedding = self.embedding_model.encode(chunk).tolist()
                            
                            # Create vector
                            vector = {
                                'id': f"{filename.replace('.docx', '')}_chunk_{i}",
                                'values': embedding,
                                'metadata': {
                                    'content': chunk,
                                    'filename': filename,
                                    'document_id': filename.replace('.docx', ''),
                                    'content_type': metadata['content_type'],
                                    'category': metadata['category'],
                                    'document_type': metadata['document_type'],
                                    'source': metadata['source'],
                                    'chunk_index': i,
                                    'is_official': True
                                }
                            }
                            vectors.append(vector)
                
                except Exception as e:
                    print(f"❌ Error processing {filename}: {str(e)}")
            else:
                print(f"❌ File not found: {filename}")
        
        # Add ADGM regulations summary (since we can't download the full PDF easily)
        regulations_content = self._get_adgm_regulations_summary()
        reg_chunks = self._chunk_content(regulations_content)
        
        for i, chunk in enumerate(reg_chunks):
            embedding = self.embedding_model.encode(chunk).tolist()
            vector = {
                'id': f"adgm_regulations_2020_chunk_{i}",
                'values': embedding,
                'metadata': {
                    'content': chunk,
                    'filename': 'adgm_companies_regulations_2020.txt',
                    'document_id': 'adgm_regulations_2020',
                    'content_type': 'official_regulation',
                    'category': 'legal_framework',
                    'document_type': 'ADGM Companies Regulations 2020',
                    'source': 'ADGM Legal Framework',
                    'chunk_index': i,
                    'is_official': True
                }
            }
            vectors.append(vector)
        
        # Upsert to Pinecone
        if vectors:
            batch_size = 100
            for i in range(0, len(vectors), batch_size):
                batch = vectors[i:i + batch_size]
                try:
                    self.index.upsert(vectors=batch)
                    print(f"Upserted batch {i//batch_size + 1}/{(len(vectors)-1)//batch_size + 1}")
                except Exception as e:
                    print(f"Error upserting batch: {str(e)}")
            
            print(f"✅ Successfully loaded {len(vectors)} official ADGM document chunks")
    
    def _read_docx_file(self, file_path: Path) -> str:
        """Read content from DOCX file"""
        try:
            doc = Document(file_path)
            content = []
            
            # Extract paragraph text
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text.strip())
            
            # Extract table text
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            content.append(cell.text.strip())
            
            return '\n'.join(content)
            
        except Exception as e:
            print(f"Error reading DOCX file {file_path}: {str(e)}")
            return ""
    
    def _chunk_content(self, content: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """Split content into overlapping chunks"""
        words = content.split()
        chunks = []
        
        for i in range(0, len(words), chunk_size - overlap):
            chunk = ' '.join(words[i:i + chunk_size])
            if chunk.strip():
                chunks.append(chunk)
        
        return chunks
    
    def analyze_document(self, document_content: str, process_type: str) -> str:
        """Analyze document using official ADGM knowledge"""
        try:
            query_embedding = self.embedding_model.encode(document_content).tolist()
            
            # Search for official ADGM content only
            search_results = self.index.query(
                vector=query_embedding,
                top_k=5,
                include_metadata=True,
                filter={'is_official': True}
            )
            
            # Extract official ADGM context
            context_chunks = []
            for match in search_results.get('matches', []):
                if match['score'] > 0.7:
                    metadata = match['metadata']
                    context_chunks.append(f"[{metadata['source']} - {metadata['document_type']}]: {metadata['content']}")
            
            context = '\n\n'.join(context_chunks[:3])
            
            # Generate analysis with official context
            analysis_prompt = f"""
            You are an expert ADGM legal assistant with access to official ADGM documents. 
            Analyze the document against official ADGM regulations and templates.
            
            Process Type: {process_type}
            
            Official ADGM Reference Context:
            {context}
            
            Document Content to Analyze:
            {document_content[:2000]}...
            
            Provide analysis based ONLY on official ADGM sources:
            1. Compliance with official ADGM templates and regulations
            2. Specific violations of ADGM Companies Regulations 2020
            3. Required corrections with exact ADGM regulatory citations
            4. Jurisdiction compliance (ADGM Courts vs UAE Federal)
            
            Base all recommendations on the official ADGM sources provided above.
            """
            
            analysis = self.gemini_client.generate_response(analysis_prompt)
            return analysis
            
        except Exception as e:
            return f"Error in official ADGM analysis: {str(e)}"
    
    def search_knowledge_base(self, query: str, top_k: int = 5) -> List[Dict]:
        """Search official ADGM knowledge base"""
        try:
            query_embedding = self.embedding_model.encode(query).tolist()
            
            results = self.index.query(
                vector=query_embedding,
                top_k=top_k,
                include_metadata=True,
                filter={'is_official': True}
            )
            
            return [
                {
                    'content': match['metadata']['content'],
                    'document_id': match['metadata']['document_id'],
                    'document_type': match['metadata']['document_type'],
                    'source': match['metadata']['source'],
                    'category': match['metadata']['category'],
                    'filename': match['metadata']['filename'],
                    'score': match['score']
                }
                for match in results.get('matches', [])
                if match['score'] > 0.6
            ]
            
        except Exception as e:
            print(f"Error searching official knowledge base: {str(e)}")
            return []
    
    def force_refresh_knowledge_base(self):
        """Force refresh knowledge base with latest official documents"""
        print("Force refreshing knowledge base...")
        
        # Delete existing vectors
        self.index.delete(delete_all=True)
        
        # Wait for deletion to complete
        time.sleep(10)
        
        # Re-download and populate
        self._download_official_documents()
        self._populate_from_official_documents()
        
        print("Knowledge base refreshed with latest official documents!")
    
    def _get_adgm_regulations_summary(self) -> str:
        """Concise summary of key ADGM regulations"""
        return """
        ADGM COMPANIES REGULATIONS 2020 - KEY COMPLIANCE PROVISIONS
        
        JURISDICTION REQUIREMENTS:
        - All companies incorporated in ADGM are subject to ADGM law and the exclusive jurisdiction of ADGM Courts
        - Any reference to UAE Federal Courts, Dubai Courts, or Abu Dhabi Courts is non-compliant
        - Documents must reference ADGM Companies Regulations 2020, not UAE Commercial Code
        
        INCORPORATION REQUIREMENTS:
        - Private companies must use ADGM Model Articles or compliant custom articles
        - No Memorandum of Association required for ADGM entities
        - Must file via ADGM Online Registry Solution only
        - Registered office must be in ADGM
        
        MANDATORY DOCUMENT ELEMENTS:
        - Proper jurisdiction clauses specifying ADGM
        - Share capital declaration for companies limited by shares
        - Directors' powers and appointment procedures
        - Proper execution with signatures and dates
        
        PROHIBITED REFERENCES:
        - UAE Federal Law or UAE Commercial Code (use ADGM Companies Regulations 2020)
        - UAE Federal Courts or Dubai Courts (use ADGM Courts)
        - UAE Ministry of Economy (use ADGM Registration Authority)
        - DIFC references (use Abu Dhabi Global Market)
        
        COMPLIANCE OBLIGATIONS:
        - Annual confirmation statements required
        - Maintain accurate beneficial ownership records
        - Keep statutory registers up to date
        - File changes within prescribed timeframes
        """
